"""
文档处理服务
支持 PDF、Word、TXT 等格式的文档解析
"""
import os
import re
from typing import List, Dict, Any
import PyPDF2
import pdfplumber
from flask import current_app


class DocumentProcessor:
    """文档处理器"""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        从 PDF 文件中提取文本
        
        Args:
            file_path: PDF 文件路径
            
        Returns:
            提取的文本内容
        """
        try:
            text = ""
            
            # 优先使用 pdfplumber（效果更好）
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                current_app.logger.warning(f'pdfplumber 解析失败，尝试使用 PyPDF2: {str(e)}')
                
                # 备用方案：使用 PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            
            # 清理文本
            text = DocumentProcessor._clean_text(text)
            
            current_app.logger.info(f'PDF 解析成功，提取文本长度: {len(text)} 字符')
            return text
            
        except Exception as e:
            current_app.logger.error(f'PDF 解析失败: {str(e)}', exc_info=True)
            raise Exception(f'PDF 文件解析失败: {str(e)}')
    
    @staticmethod
    def extract_text_from_word(file_path: str) -> str:
        """
        从 Word 文件中提取文本
        
        Args:
            file_path: Word 文件路径
            
        Returns:
            提取的文本内容
        """
        try:
            # 尝试导入 python-docx
            try:
                import docx
            except ImportError:
                raise Exception('需要安装 python-docx: pip install python-docx')
            
            doc = docx.Document(file_path)
            text = ""
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            # 清理文本
            text = DocumentProcessor._clean_text(text)
            
            current_app.logger.info(f'Word 解析成功，提取文本长度: {len(text)} 字符')
            return text
            
        except Exception as e:
            current_app.logger.error(f'Word 解析失败: {str(e)}', exc_info=True)
            raise Exception(f'Word 文件解析失败: {str(e)}')
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """
        从 TXT 文件中提取文本
        
        Args:
            file_path: TXT 文件路径
            
        Returns:
            提取的文本内容
        """
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    
                    # 清理文本
                    text = DocumentProcessor._clean_text(text)
                    
                    current_app.logger.info(
                        f'TXT 解析成功（编码: {encoding}），提取文本长度: {len(text)} 字符'
                    )
                    return text
                    
                except UnicodeDecodeError:
                    continue
            
            raise Exception('无法识别文件编码，请确保文件为 UTF-8 或 GBK 编码')
            
        except Exception as e:
            current_app.logger.error(f'TXT 解析失败: {str(e)}', exc_info=True)
            raise Exception(f'TXT 文件解析失败: {str(e)}')
    
    @staticmethod
    def extract_text(file_path: str, file_type: str = None) -> str:
        """
        根据文件类型提取文本
        
        Args:
            file_path: 文件路径
            file_type: 文件类型（可选，自动从扩展名判断）
            
        Returns:
            提取的文本内容
        """
        if not file_type:
            _, ext = os.path.splitext(file_path)
            file_type = ext.lower().lstrip('.')
        
        if file_type == 'pdf':
            return DocumentProcessor.extract_text_from_pdf(file_path)
        elif file_type in ['doc', 'docx']:
            return DocumentProcessor.extract_text_from_word(file_path)
        elif file_type == 'txt':
            return DocumentProcessor.extract_text_from_txt(file_path)
        else:
            raise Exception(f'不支持的文件类型: {file_type}')
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """
        清理文本
        
        Args:
            text: 原始文本
            
        Returns:
            清理后的文本
        """
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        
        # 移除特殊字符（保留中文、英文、数字、常用标点）
        text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s\.,!?;:，。！？；：、""''（）\(\)\-—]', '', text)
        
        # 移除多余的换行
        text = re.sub(r'\n\s*\n', '\n', text)
        
        # 去除首尾空白
        text = text.strip()
        
        return text
    
    @staticmethod
    def split_text(
        text: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        separators: List[str] = None
    ) -> List[str]:
        """
        智能分块文本
        
        Args:
            text: 要分块的文本
            chunk_size: 每个块的大小（字符数）
            chunk_overlap: 块之间的重叠大小
            separators: 分隔符列表（按优先级）
            
        Returns:
            分块后的文本列表
        """
        if not separators:
            separators = ['\n\n', '\n', '。', '！', '？', ';', '；', ',', '，', ' ']
        
        chunks = []
        current_chunk = ""
        
        # 先按段落分割
        paragraphs = text.split('\n')
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            # 如果当前块加上新段落不超过限制，直接添加
            if len(current_chunk) + len(paragraph) <= chunk_size:
                if current_chunk:
                    current_chunk += "\n" + paragraph
                else:
                    current_chunk = paragraph
            else:
                # 当前块已满，保存并开始新块
                if current_chunk:
                    chunks.append(current_chunk)
                    
                    # 添加重叠部分
                    if chunk_overlap > 0:
                        overlap_text = current_chunk[-chunk_overlap:]
                        current_chunk = overlap_text + "\n" + paragraph
                    else:
                        current_chunk = paragraph
                else:
                    current_chunk = paragraph
                
                # 如果单个段落就超过了块大小，需要进一步分割
                while len(current_chunk) > chunk_size:
                    # 找到合适的分割点
                    split_point = chunk_size
                    for sep in separators:
                        pos = current_chunk.rfind(sep, 0, chunk_size)
                        if pos > chunk_size // 2:  # 至少要分割一半以上
                            split_point = pos + len(sep)
                            break
                    
                    # 保存当前块
                    chunks.append(current_chunk[:split_point].strip())
                    
                    # 准备下一块（带重叠）
                    if chunk_overlap > 0:
                        overlap_start = max(0, split_point - chunk_overlap)
                        current_chunk = current_chunk[overlap_start:]
                    else:
                        current_chunk = current_chunk[split_point:]
        
        # 添加最后一块
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        current_app.logger.info(f'文本分块完成: {len(chunks)} 个块')
        return chunks
    
    @staticmethod
    def create_chunks_with_metadata(
        text: str,
        document_id: str,
        document_name: str,
        kb_id: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ) -> List[Dict[str, Any]]:
        """
        创建带元数据的文本块
        
        Args:
            text: 要分块的文本
            document_id: 文档ID
            document_name: 文档名称
            kb_id: 知识库ID
            chunk_size: 块大小
            chunk_overlap: 重叠大小
            
        Returns:
            包含元数据的块列表
        """
        # 分块
        chunks = DocumentProcessor.split_text(text, chunk_size, chunk_overlap)
        
        # 添加元数据
        chunks_with_metadata = []
        for i, chunk_text in enumerate(chunks):
            chunk_id = f"{document_id}_chunk_{i}"
            chunks_with_metadata.append({
                'id': chunk_id,
                'content': chunk_text,
                'metadata': {
                    'document_id': document_id,
                    'document_name': document_name,
                    'kb_id': kb_id,
                    'chunk_index': i,
                    'chunk_total': len(chunks),
                    'source': document_name
                }
            })
        
        return chunks_with_metadata


# 全局文档处理器实例
document_processor = DocumentProcessor()

